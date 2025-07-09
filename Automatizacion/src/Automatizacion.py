#LibreOffice Calc Automation Script
# -*- coding: utf-8 -*-
#* This script automates the process of filling a Word document template with employee data from a JSON file.
import json
import re 
import polars as pl
import json
from docx import Document
from docx.shared import Pt



# Please read the readme
#if u have other document you need change the name in this line.
df = pl.read_excel('../data/Formato_Capacitacion.xlsx') #Here

empleados = []
for row in df.to_dicts():
    empleado = {
        "datosDelEmpleado": {
            "nombre": row["Nombre del Empleado"],
            "curp": row["CURP"],
            "ocupacion": row["Ocupacion"],
            "puesto": row["Puesto"]
        },
        "DatosEmpresa": {
            "nombreEmpresa": row["Nombre de la Empresa"],
            "SHCP": row["Registro SHCP (RFC)"]
        },
        "DatosDelProgramaDeCapacitacion": {
            "nombreDelPrograma": row["Nombre del Programa"],
            "duracion": row["Duración del Programa"],
            "perido": {
                "fechaInicio": row["Fecha de Inicio"],
                "fechaFin": row["Fecha de Fin"]
            },
            "areaTematicaDelCurso": row["Area Temática del Curso"],
            "nombreDelAgenteCapacitador": row["Nombre del Agente Capacitador"]
        }
    }
    empleados.append(empleado)

# save the resul in a file .json
with open('empleados.json', 'w', encoding='utf-8') as f:
    json.dump(empleados, f, ensure_ascii=False, indent=4)


def cleaner_string(string):
    result =  []
    for n in string:
        if not n == "-":
            result.append(n)
    return result

#Load the file empleados.json 
with open('empleados.json', 'r', encoding='utf-8') as f:
    empleados = json.load(f)

listEmpleados = len(empleados)
doc = Document('Plantilla_.docx')

for empleado in range(1, listEmpleados + 1):
    print(f"Procesando empleado {empleado} de {listEmpleados}")

    #Declarate var
    #data of employe
    nombre_empleado =  empleados[empleado-1]["datosDelEmpleado"]["nombre"]
    curp_empleado  =   empleados[empleado-1]["datosDelEmpleado"]["curp"]
    ocupacion_empleado = empleados[empleado-1]["datosDelEmpleado"]["ocupacion"]
    puest_empleado  =  empleados[empleado-1]["datosDelEmpleado"]["puesto"]
    #data of empresa
    nombre_empresa =  empleados[empleado-1]["DatosEmpresa"]["nombreEmpresa"]
    SHCP = empleados[empleado-1]["DatosEmpresa"]["SHCP"]
    #data of program the curse
    nombre_programa =  empleados[empleado-1]["DatosDelProgramaDeCapacitacion"]["nombreDelPrograma"]
    duracion = empleados[empleado-1]["DatosDelProgramaDeCapacitacion"]["duracion"]
    fecha_inicio_dirty = empleados[empleado-1]["DatosDelProgramaDeCapacitacion"]["perido"]["fechaInicio"]
    fecha_fin_dirty = empleados[empleado-1]["DatosDelProgramaDeCapacitacion"]["perido"]["fechaFin"]
    area_tematica_del_curso =   empleados[empleado-1]["DatosDelProgramaDeCapacitacion"]["areaTematicaDelCurso"]
    nombre_capacitador =   empleados[empleado-1]["DatosDelProgramaDeCapacitacion"]["nombreDelAgenteCapacitador"]
    fecha_inicio = cleaner_string(fecha_inicio_dirty)
    fecha_fin = cleaner_string(fecha_fin_dirty)

    #* Here create the automiation for the frist table


    tableDataEmpleado = doc.tables[1]
    #set name employe
    cell_name = tableDataEmpleado.cell(2, 0)
    cell_name.text = ""
    parrafo_name = cell_name.paragraphs[0]
    run_name = parrafo_name.add_run(nombre_empleado)
    run_name.font.name = 'Arial'
    run_name.font.size = Pt(9)
    #set curp
    position  =  0
    for letter in range(1,19):

        cell_curp =  tableDataEmpleado.cell(4,letter -1)
        cell_curp.text = ''
        parrafo_curp = cell_curp.paragraphs[0]
        run_curp  = parrafo_curp.add_run(curp_empleado[position])
        run_curp.font.name =  'Arial'
        run_curp.font.size =  Pt(9)
        position+=1

    #set ocupation
    cell_Ocupation =  tableDataEmpleado.cell(4,18)
    cell_Ocupation.text =  ""
    parrafo_ocupation =  cell_Ocupation.paragraphs[0]
    run_Ocupation = parrafo_ocupation.add_run(ocupacion_empleado)
    run_Ocupation.font.name =  'Arial'
    run_Ocupation.font.size =  Pt(9)
    #set position
    cell_position = tableDataEmpleado.cell(6,1)
    cell_position.text = ""
    parrafo_position = cell_position.paragraphs[0]
    run_position = parrafo_position.add_run(puest_empleado)
    run_position.font.name = 'Arial'
    run_position.font.size = Pt(9)
    

    #* Here create the automiation for the second table
    tableDatosEmmpresa = doc.tables[2]

    #set name empresa
    cell_name_empresa =  tableDatosEmmpresa.cell(2, 0)
    cell_name_empresa.text = ""
    parrafo_name_empresa =  cell_name_empresa.paragraphs[0]
    run_name_empresa =  parrafo_name_empresa.add_run(nombre_empresa)
    run_name_empresa.font.name =  'Arial'
    run_name_empresa.font.size =  Pt(9)
    # set SHCP
    position_spch = 0
    for caracter in range(1,15):
        cell_shpc =  tableDatosEmmpresa.cell(4, caracter-1)
        cell_shpc.text = ''
        parrafo_shpc =  cell_shpc.paragraphs[0]
        run_shpc =  parrafo_shpc.add_run(SHCP[position_spch])
        run_shpc.font.name = 'Arial'
        run_shpc.font.size = Pt(9)
        position_spch +=1

    #* Here create the automiation for the thrid table
    tableDataDelPrograma = doc.tables[3]

    #set name curse
    cell_name_curse =  tableDataDelPrograma.cell(2,0)
    cell_name_curse.text =  ""
    parrafo_name_curse =  cell_name_curse.paragraphs[0]
    run_name_curse = parrafo_name_curse.add_run(nombre_programa)
    run_name_curse.font.name = 'Arial'
    run_name_curse.font.size =  Pt(8)

    #set duration curse
    cell_duration_curse  =  tableDataDelPrograma.cell(4,0)
    cell_duration_curse.text = ""
    parrafo_duration =  cell_duration_curse.paragraphs[0]
    run_duration  = parrafo_duration.add_run(duracion)
    run_duration.font.name = 'Arial'
    run_duration.font.size =  Pt(9)
    #set dates
    position_date_start = 0
    for date_caracter in range(4,12):
        cell_date_start =  tableDataDelPrograma.cell(4,date_caracter-1)
        cell_date_start.text = ''
        parrafo_date_inicio  =  cell_date_start.paragraphs[0]
        parrafo_date_inicio =  parrafo_date_inicio.add_run(fecha_inicio[position_date_start])
        parrafo_date_inicio.font.name = 'Arial'
        parrafo_date_inicio.font.size =  Pt(9)
        position_date_start+=1
    position_date_fin = 0
    for date_caracter in range(13,21):
        cell_date_start = tableDataDelPrograma.cell(4, date_caracter - 1)
        cell_date_start.text = ''
        parrafo_date_inicio = cell_date_start.paragraphs[0]
        parrafo_date_inicio = parrafo_date_inicio.add_run(fecha_inicio[position_date_fin])
        parrafo_date_inicio.font.name = 'Arial'
        parrafo_date_inicio.font.size = Pt(9)
        position_date_fin += 1
    #set area tematic curse
    cell_area_curse  =  tableDataDelPrograma.cell(6,0)
    cell_area_curse.text = ''
    parrafo_area_curse = cell_area_curse.paragraphs[0]
    run_are_curse =  parrafo_area_curse.add_run(area_tematica_del_curso)
    run_are_curse.font.name = 'Arial'
    run_are_curse.font.size = Pt(8)

    #set name capacitator
    cell_name_capacitator = tableDataDelPrograma.cell(8,0)
    cell_name_capacitator.text = ''
    parrafo_name_capacitator = cell_name_capacitator.paragraphs[0]
    run_name_capacitator = parrafo_name_capacitator.add_run(nombre_capacitador)
    run_name_capacitator.font.name = 'Arial'
    run_name_capacitator.font.size = Pt(8)

    #Save Files in the carpt asigned ("./Save_Files")
    nombre_archivo = re.sub(r'[\\/*?:"<>|]', "", empleados[empleado-1]["datosDelEmpleado"]["nombre"])
    doc.save(f'Save_Files/empleado_{nombre_archivo}.docx')


